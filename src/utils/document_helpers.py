from docx import Document
from docx.shared import Inches
import os
from datetime import datetime
import logging
from typing import Union, Dict
from utils.image_helpers import generate_and_save_images

logger = logging.getLogger(__name__)

def display_explanation(explanation_dict: Dict, user_intent: str):
    """Display the parsed explanation and generate document."""
    logger.info("Starting explanation display")
    
    try:
        doc = Document()
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        sanitized_intent = "".join(c for c in user_intent if c.isalnum() or c in (' ', '-', '_')).strip()
        sanitized_intent = sanitized_intent.replace(' ', '_').lower()
        output_folder = os.path.join("output", f"{timestamp}_{sanitized_intent}")
        os.makedirs(output_folder, exist_ok=True)
        
        doc.add_heading(explanation_dict['title'], 0)
        doc.add_paragraph(explanation_dict['introduction'])
        
        images_folder = generate_and_save_images(explanation_dict, user_intent)
        logger.info(f"Images saved in: {images_folder}")
        
        for step in explanation_dict['steps']:
            doc.add_heading(f"{step['heading']}", level=1)
            doc.add_paragraph(step['text'])
            
            image_path = os.path.join(images_folder, f"step_{step['step_number']}.png")
            if os.path.exists(image_path):
                desc_para = doc.add_paragraph()
                desc_para.add_run(step['image_description'])
                doc.add_picture(image_path, width=Inches(6))

            if 'transition' in step:
                transition_para = doc.add_paragraph()
                transition_para.add_run(step['transition'])
        
        doc.add_heading('Conclusion', level=1)
        doc.add_paragraph(explanation_dict['conclusion'])
        
        doc_path = os.path.join(output_folder, f"{sanitized_intent}.docx")
        doc.save(doc_path)
        logger.info(f"Document saved as: {doc_path}")
        
        return output_folder
        
    except Exception as e:
        logger.error("Error creating explanation document", exc_info=True)
        raise 