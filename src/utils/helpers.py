import os
import logging
from datetime import datetime
from typing import Dict, Union
import requests
from openai import OpenAI
from config.config_manager import ConfigManager
import yaml
from docx import Document
from docx.shared import Inches, Pt
import sys

# Set up logger
logger = logging.getLogger(__name__)

# Get config instance
config = ConfigManager()

# Initialize OpenAI client
client = OpenAI(
    api_key=config.get('openai.api_key')
)

def setup_logging():
    """Configure logging to output to both file and console."""
    # Create logger
    logger = logging.getLogger()
    logger.setLevel(logging.DEBUG)
    
    # Create formatters
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    
    # Create console handler
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setFormatter(formatter)
    console_handler.setLevel(logging.DEBUG)
    
    # Create file handler
    file_handler = logging.FileHandler('debug.log')
    file_handler.setFormatter(formatter)
    file_handler.setLevel(logging.DEBUG)
    
    # Add handlers to logger
    logger.addHandler(console_handler)
    logger.addHandler(file_handler)

# Call this function at the start of your application
setup_logging()

def load_system_prompt(file_path: str) -> str:
    """Load system prompt from file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def clean_yaml_string(yaml_str: str) -> str:
    """Clean up YAML string by removing code block markers and normalizing content."""
    # Split into lines for better processing
    lines = yaml_str.split('\n')
    
    # Remove code block markers at start and end
    if lines and any(lines[0].startswith('```')):
        lines = lines[1:]
    if lines and any(lines[-1].startswith('```')):
        lines = lines[:-1]
    
    # Join lines back together and clean up whitespace
    cleaned_str = '\n'.join(lines).strip()
    
    # Replace problematic characters and normalize quotes
    cleaned_str = (cleaned_str
        .replace('\u2018', "'")  # Replace fancy single quotes
        .replace('\u2019', "'")
        .replace('\u201c', '"')  # Replace fancy double quotes
        .replace('\u201d', '"')
        .replace('\u2013', '-')  # Replace en-dash
        .replace('\u2014', '--')  # Replace em-dash
    )
    
    return cleaned_str

def parse_yaml_response(yaml_str: str) -> Union[Dict, str]:
    """Parse YAML string to dictionary."""
    try:
        # Remove code block markers if they exist
        cleaned_str = yaml_str
        if cleaned_str.startswith('```yaml'):
            cleaned_str = cleaned_str[7:]  # Remove ```yaml
        elif cleaned_str.startswith('```'):
            cleaned_str = cleaned_str[3:]  # Remove ```
        if cleaned_str.endswith('```'):
            cleaned_str = cleaned_str[:-3]  # Remove trailing ```
        
        # Strip any extra whitespace
        cleaned_str = cleaned_str.strip()
        # Replace problematic characters and normalize quotes
        cleaned_str = (cleaned_str
        .replace('\u2018', "'")  # Replace fancy single quotes
        .replace('\u2019', "'")
        .replace('\u201c', '"')  # Replace fancy double quotes
        .replace('\u201d', '"')
        .replace('\u2013', '-')  # Replace en-dash
        .replace('\u2014', '--')  # Replace em-dash
        )
        return yaml.safe_load(cleaned_str)
    except yaml.YAMLError as e:
        logger.error(f"Error parsing YAML: {e}\nInput that caused error:\n{'-' * 50}\n{cleaned_str}\n{'-' * 50}")
        return yaml_str

def get_completion(user_intent: str) -> str:


    # Load and format system prompt
    system_prompt = load_system_prompt(config.get('openai.system_prompt_path'))
    formatted_prompt = system_prompt.format(user_intent=user_intent)

    # Create completion
    completion = client.chat.completions.create(
        model=config.get('openai.model'),
        messages=[
            {
                "role": "system", 
                "content": formatted_prompt
            }
        ]
    )

    return completion.choices[0].message.content

def generate_and_save_images(explanation_dict: Dict, user_intent: str):
    """Generate DALL-E images for each step and save them in a dedicated folder."""
    logger.info(f"Starting image generation for concept: {user_intent}")
    folder_name = os.path.join("images", f"{user_intent}")

    
    try:
        os.makedirs(folder_name, exist_ok=True)
        logger.debug(f"Created directory: {folder_name}")
        
        config = ConfigManager()
        client = OpenAI(api_key=config.get('openai.api_key'))
        
        for step in explanation_dict['steps']:
            try:
                logger.info(f"Generating image for step {step['step_number']}")
                logger.debug(f"Image description: {step['image_description']}")
                
                response = client.images.generate(
                    model="dall-e-3",
                    prompt=step['image_description'],
                    size="1024x1024",
                    quality="standard",
                    n=1,
                )
                
                image_url = response.data[0].url
                logger.debug(f"Received image URL for step {step['step_number']}")
                
                image_data = requests.get(image_url).content
                image_path = f"{folder_name}/step_{step['step_number']}.png"
                
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                logger.info(f"Saved image for step {step['step_number']}")
                
            except Exception as e:
                logger.error(f"Error generating image for step {step['step_number']}: {e}", exc_info=True)
                
        logger.info(f"Completed image generation. Images saved in: {folder_name}")
        return folder_name
        
    except Exception as e:
        logger.error(f"Error in image generation process: {e}", exc_info=True)
        raise

def display_explanation(explanation_dict: Dict, user_intent: str):
    """Display the parsed explanation and generate images and Word document."""
    logger.info("Starting explanation display")
    
    try:
        # Create Word document
        doc = Document()
        
        # Create output folder
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        sanitized_intent = "".join(c for c in user_intent if c.isalnum() or c in (' ', '-', '_')).strip()
        sanitized_intent = sanitized_intent.replace(' ', '_').lower()
        output_folder = os.path.join("output", f"{timestamp}_{sanitized_intent}")
        os.makedirs(output_folder, exist_ok=True)
        
        # Add title
        doc.add_heading(explanation_dict['title'], 0)
        
        # Add introduction
        doc.add_paragraph(explanation_dict['introduction'])
        
        # Generate images and add steps
        images_folder = generate_and_save_images(explanation_dict, user_intent)
        logger.info(f"Images saved in: {images_folder}")
        
        for step in explanation_dict['steps']:
            # Add step heading
            doc.add_heading(f"{step['heading']}", level=1)
            
            # Add step text
            doc.add_paragraph(step['text'])
            
            # Add image if it exists
            image_path = os.path.join(images_folder, f"step_{step['step_number']}.png")
            if os.path.exists(image_path):
                # Add image description before
                desc_para = doc.add_paragraph()
                desc_para.add_run(step['image_description'])
                
                # Add the image
                doc.add_picture(image_path, width=Inches(6))

            # Add transition if it exists
            if 'transition' in step:
                transition_para = doc.add_paragraph()
                transition_para.add_run(step['transition'])
        
        # Add conclusion
        doc.add_heading('Conclusion', level=1)
        doc.add_paragraph(explanation_dict['conclusion'])
        
        # Save the document
        doc_path = os.path.join(output_folder, f"{sanitized_intent}.docx")
        doc.save(doc_path)
        logger.info(f"Document saved as: {doc_path}")
        
        return output_folder
        
    except Exception as e:
        logger.error("Error creating explanation document", exc_info=True)
        raise 